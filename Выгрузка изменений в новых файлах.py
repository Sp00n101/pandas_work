import pandas as pd
import tkinter as tk
import os
import openpyxl
import xlrd
from docx import Document
import win32com.client
import re
from openpyxl.workbook.child import INVALID_TITLE_REGEX


def replace_rus_to_eng(s):
    rus = "уУеЕНхХаАрРоОсСТМК"
    end = "yYeEHxXaApPoOcCTMK"
    for i in range(len(rus)):
        s = s.replace(rus[i], end[i])
    return s


def abs_path(dir):
    fail = os.listdir(dir)
    ab_fail = []
    for i in range(len(fail)):
        ab_fail.append(dir+fail[i])
    return ab_fail

def excel_read(path, name):
    excel = []
    for i in path:
        print(i)
        if any(replace_rus_to_eng(sub) in replace_rus_to_eng(i) for sub in name) and (".XLSM" in i or ".XLSX" in i or ".xlsx" in i or ".xlsm" in i):

            try:
                wb = openpyxl.load_workbook(i)
            except:
                excel.append({i: ["Файл был пропущен"]})
                continue
            highlighted_cells_sheet = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]

                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.fill.start_color.index == 'FFFFFF00':  # Жёлтый цвет заливки (FFFFFF00 - hex код для жёлтого)
                            highlighted_cells_sheet.append(cell.value)

            excel.append({i.replace(path_dock, ""): highlighted_cells_sheet})
        elif any(replace_rus_to_eng(sub) in replace_rus_to_eng(i) for sub in name) and (".XLS" in i or ".xls" in i):
            try:
                workbook = xlrd.open_workbook(i, formatting_info=True)
            except:
                excel.append({i: ["Файл был пропущен"]})
                continue

            yellow_cells = []
            for sheet in workbook.sheets():
                # Проходим по всем ячейкам и добавляем значения жёлтых ячеек в массив

                for row in range(sheet.nrows):
                    for col in range(sheet.ncols):
                        cell = sheet.cell(row, col)
                        # Получаем цвет фона ячейки
                        cell_color_index = cell.xf_index
                        if cell_color_index is not None:
                            cell_color = workbook.xf_list[cell_color_index].background.pattern_colour_index
                            # Проверяем, является ли цвет фона жёлтым (цвет для жёлтого - 13)
                            if cell_color == 13:
                                yellow_cells.append(cell.value)
            excel.append({i.replace(path_dock, ""): yellow_cells})
        elif any(replace_rus_to_eng(sub) in replace_rus_to_eng(i) for sub in name) and (".DOCX" in i or ".docx" in i):
            doc = Document(i)

            yellow_highlighted_text = []

            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run.text and run.font.highlight_color == 7:
                        yellow_highlighted_text.append(run.text)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if run.text and run.font.highlight_color == 7:
                                    yellow_highlighted_text.append(run.text)
            excel.append({i.replace(path_dock, ""): yellow_highlighted_text})
            print("OK")
        elif any(replace_rus_to_eng(sub) in replace_rus_to_eng(i) for sub in name) and (".doc" in i):
            word = win32com.client.Dispatch("Word.Application")
            print(1)
            doc = word.Documents.Open(i)
            print(2)
            highlighted_text = []
            print(3)
            for paragraph in doc.Paragraphs:
                for run in paragraph.Range.Words:
                    if run.HighlightColorIndex == 7:  # Жёлтый цвет заливки (7 - код для жёлтого)
                        highlighted_text.append(run.Text)

            excel.append({i.replace(path_dock, ""): highlighted_text})
            print(8)
            doc.Close()
            word.Quit()
    return excel

def fail_to_read(path):
    df = pd.read_excel(path, sheet_name=sheet_name, skiprows=72, usecols="A:D,K:L")
    df = df.dropna(axis=1, how='all')
    df.columns = df.iloc[0]
    df = df.drop(index=0)
    df = df[df['PD'].str.contains('M') | df['PDC'].str.contains('M') | df['PD'].str.contains('М') | df['PDC'].str.contains('М')]

    fail_read = list(df["Новый № документа"])
    return fail_read

def on_button_click():
    global path_dock
    path_dock = entry.get()
    root.destroy()

def tkin():
    global root
    root = tk.Tk()
    root.title("Ввод данных")

    lbl = tk.Label(root, text="Введите путь до файлов:")
    lbl.pack()

    # Поле ввода
    global entry
    entry = tk.Entry(root)
    entry.pack(padx=20, pady=10)

    # Кнопка
    button = tk.Button(root, text="Сохранить и закрыть", command=on_button_click)
    button.pack(pady=10)

    # Запускаем цикл обработки событий
    root.mainloop()

path_dock = ''

files_in_directory = os.listdir("./")
file_name = "нет пути"
sheet_name = "1ая страница"

for file in files_in_directory:
    if ".xlsm" in file and "062_F_R_VKMB" in file:
        file_name = "./" + file
        break

if file_name != '':
    path_dock = "C:\PycharmProjects\Izmeneniy\direct\\"

    name_dock = fail_to_read(file_name)
    path = abs_path(path_dock)
    data = excel_read(path, name_dock)
    dfs = []
    for d in data:
        key = list(d.keys())[0]
        values = d[key]
        if values == []:
            values = ["Никаких изменений не найденно"]
        df = pd.DataFrame({'File': [key] * len(values), 'Values': values})
        dfs.append(df)

    result_df = pd.concat(dfs, ignore_index=True)
    result_df['Values'] = result_df['Values'].astype(str)
    result_df = result_df.groupby('File')['Values'].apply(lambda x: ', '.join(x)).reset_index()
    with pd.ExcelWriter("./Данные_за_месяц.xlsx") as writer:
        result_df.to_excel(writer, sheet_name=re.sub(INVALID_TITLE_REGEX, '_', "Изменения"))